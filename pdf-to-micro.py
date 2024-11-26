import pdfplumber
import tkinter as tk
import os
import re  

from openpyxl import Workbook
from openpyxl.styles import Font
from docx import Document
from PyPDF2 import PdfReader
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image
import io

# ================== SUB-FUNCTIONS ================== #
# XLSX FIND HEADER
def find_header(page, table_bbox, table_settings):
    # Define the bounding box for the header
    h_top = max(table_bbox[1] - 35, 0)  # Top boundary of header bbox
    h_x0 = max(table_bbox[0] - 15, 0)  # Left boundary of header bbox
    header_bbox = (h_x0, h_top, table_bbox[2], table_bbox[1])
    x_tol = table_settings["text_x_tolerance"]
    y_tol = table_settings["text_y_tolerance"]
    # Extract words within the bounding box
    words = [
        word["text"] for word in page.extract_words(x_tolerance=x_tol, y_tolerance=y_tol)
        if word["x0"] >= header_bbox[0] and
           word["x1"] <= header_bbox[2] and
           word["top"] >= header_bbox[1] and
           word["bottom"] <= header_bbox[3]
    ]

    # Combine words into a single string
    header = " ".join(words)

    # Return the extracted header text
    return header.strip()



# ================== MAIN FUNCTIONS ================== #

# PDF TO XLSX
def pdf_to_xlsx(pdf_path, xlsx_path, user_settings=None, user_file_settings=None):
    wb = Workbook()
    ws = wb.active

    # if the xlsx file at the filepath exists, make sure its closed before editing
    if os.path.exists(xlsx_path):
        try:
            # Try to open the file in exclusive mode
            with open(xlsx_path, 'r+'):
                pass
        except IOError:
            # print error statement in red
            print(f"\033[91mError: The file {xlsx_path} is open. Please close it and try again.\033[0m")
            return
    
    # Ensure the debugging directory exists
    os.makedirs('debugging', exist_ok=True)
    # get the debugging directory path
    debugging_dir = os.path.abspath('debugging')

    # pull the filename from the path without the extension
    filename = os.path.splitext(os.path.basename(pdf_path))[0]

    # set init. conditions and flags
    table_continued = False
    tables_sum = 0
    page_num = 0
    pages_list = []


    # set the extract file and table settings from user input
    file_settings = {
        "include_header": True,
        "split_pages_horiz": False
    }

    table_settings = {
        "vertical_strategy": "lines",       # "lines", "lines_strict", "text", "explicit"
        "horizontal_strategy": "lines",     # "lines", "lines_strict", "text", "explicit"
        "explicit_vertical_lines": [],      # list of x-coordinates to force vertical lines
        "explicit_horizontal_lines": [],    # list of y-coordinates to force horizontal lines
        "snap_tolerance": 3,                # parallel lines within snap_tolerance points will be
        "snap_x_tolerance": 3,              #    "snapped" to the same horizontal or vertical position.
        "snap_y_tolerance": 3,
        "join_tolerance": 3,                # line segments on same infinite line within join_tolerance
        "join_x_tolerance": 3,              #     will be joined together
        "join_y_tolerance": 3,
        "edge_min_length": 3,               # lines shorter than this will be discarded
        "min_words_vertical": 3,            # for vert_strat == "text": min. words must share alignment
        "min_words_horizontal": 1,          # for horiz_strat == "text": min. words must share alignment
        "intersection_tolerance": 3,        # when combo edges into cells
        "intersection_x_tolerance": 3,      #    the intersection must be within this tolerance
        "intersection_y_tolerance": 3,
        "text_tolerance": 3,                # when searching for words, indiv letters must be
        "text_x_tolerance": 3,              #    text_tolerance apart
        "text_y_tolerance": 3,
    }
    if user_settings:
        table_settings.update(user_settings) 
    if user_file_settings:  
        file_settings.update(user_file_settings)
        print(file_settings)

    # open the pdf file
    with pdfplumber.open(pdf_path) as pdf:
        # loop through each page in the pdf
        ws.append([filename])
        for cell in ws[ws.max_row]:
            cell.font = Font(bold=True, size=18)
        ws.append([])

        if file_settings["split_pages_horiz"]:
            for page in pdf.pages:
                width = page.width
                height = page.height

                if width > height:
                    pages_list.append(page.crop((0, 0, width/2, height)))
                    pages_list.append(page.crop((width/2, 0, width, height)))
                else:
                    pages_list.append(page)
        else:
            pages_list = pdf.pages


        for page in pages_list:

            page_num += 1

            # # create image for visual debugging    
            # im = page.to_image(300)
            # im.debug_tablefinder(table_settings)
            # # save to the debugging directory
            # im.save(f"{debugging_dir}/page_{page.page_number}.png")

            
            # find all tables in the page
            tables = page.find_tables(table_settings)
            legitimate_tables = []
            # filter out tables with only 1 row or 1 column
            for table in tables:
                if len(table.columns) > 1 and len(table.rows) > 1:
                    legitimate_tables.append(table)
            
            # loop through each legitimate table
            for i, table in enumerate(legitimate_tables):
                # extract the table's data
                table_valid = False
                data = table.extract(x_tolerance=table_settings["text_x_tolerance"], y_tolerance=table_settings["text_y_tolerance"])
                # check the table isn't empty
                for row in data:
                    for cell in row:
                        if cell is not None and cell != "":
                            table_valid = True

                # for valid, non-null tables with more than 1 row & 1 column
                if table_valid:
                    if not table_continued:
                        tables_sum += 1

                        if file_settings["include_header"]:
                            # pull it's header
                            header = find_header(page, table.bbox, table_settings)
                            # add the header to the worksheet
                            ws.append([header])
                            for cell in ws[ws.max_row]:
                                cell.font = Font(bold=True)
                        ws.append([f"Table {tables_sum}, Page {page_num}"])
                        for cell in ws[ws.max_row]:
                            cell.font = Font(italic=True)

                    # add each row to the worksheet
                    for i, row in enumerate(data):
                        ws.append(row)
                        if i == 0 and not table_continued:
                            for cell in ws[ws.max_row]:
                                cell.font = Font(bold=True)

                # check if the table continues on the next page
                if i == len(legitimate_tables) - 1 and len(data) > 0 and data[-1] == ['']:
                    table_continued = True
                else:
                    table_continued = False
                    ws.append([])
                    ws.append([])
      
    wb.save(xlsx_path)
    print(f"\033[45mExcel file saved to {xlsx_path}\033[0m")



# # PDF TO DOCX
# def pdf_to_docx(pdf_path, docx_path):
#     doc = Document()
    
#     # Open PDF using pdfplumber
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             # Extract text
#             text = page.extract_text()
#             if text:
#                 paragraph = doc.add_paragraph()
#                 paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#                 run = paragraph.add_run(text)
#                 run.font.size = Pt(11)  # Adjust font size as needed
            
#             for image_obj in page.images:
#                 try:
#                     # Extract image data from PDFStream
#                     image_data = image_obj["stream"].get_data()
                    
#                     # Check if the image data is valid
#                     if image_data:
#                         # Try to open the image with PIL
#                         try:
#                             img = Image.open(io.BytesIO(image_data))
#                         except Image.UnidentifiedImageError:
#                             # If the image format is not recognized, try converting it
#                             img = Image.open(io.BytesIO(image_data)).convert("RGB")
                        
#                         # Save image to DOCX
#                         img_bytes = io.BytesIO()
#                         img.save(img_bytes, format="PNG")
#                         img_bytes.seek(0)
#                         doc.add_picture(img_bytes)
#                     else:
#                         print("No image data found.")
#                 except Exception as e:
#                     print(f"\033[91mError processing image: {e}\033[0m")
            
#             # Add a page break after each PDF page
#             doc.add_page_break()
    
#     # Save the resulting DOCX
#     doc.save(docx_path)


# ================== GUI ================== #

# Create the main window
root = tk.Tk()
root.title("PDF to Microsoft Office Converter")
# 16:10 aspect ratio
root.geometry("1000x620")
# make the window resizable
root.resizable(True, True)
# set the background color
root.configure(bg="white")
# start the main loop
root.mainloop()



































# ================== RUN TESTS ================== #

# # get path to this python file
# path = os.path.dirname(os.path.abspath(__file__))


# # VARIABLE TO STORE TABLE SETTINGS TO PASS INTO MAIN FUNCTIONS
# # Hamilton Arc Modbus OPC Settings
# ham_table_settings = {
#     "vertical_strategy": "text", 
#     # "horizontal_strategy": "lines",
#     # "explicit_vertical_lines": [],
#     # "explicit_horizontal_lines": [],
#     # "snap_tolerance": 3,
#     # "snap_x_tolerance": 3,
#     # "snap_y_tolerance": 3,
#     "join_tolerance": 10,
#     # "join_x_tolerance": 3,
#     # "join_y_tolerance": 3,
#     # "edge_min_length": 3,
#     "min_words_vertical": 3,
#     # "min_words_horizontal": 1,
#     # "intersection_tolerance": 3,
#     # "intersection_x_tolerance": 3,
#     # "intersection_y_tolerance": 3,
#     # "text_tolerance": 3,
#     # "text_x_tolerance": 3,
#     # "text_y_tolerance": 3,
# }

# ham_file_settings = {
#     "include_header": True,
#     "split_pages_horiz": True
# }

# fut_table_settings = {
#     # "vertical_strategy": "lines", 
#     # "horizontal_strategy": "lines",
#     # "explicit_vertical_lines": [],
#     # "explicit_horizontal_lines": [],
#     # "snap_tolerance": 3,
#     # "snap_x_tolerance": 3,
#     # "snap_y_tolerance": 3,
#     # "join_tolerance": 3,
#     # "join_x_tolerance": 3,
#     # "join_y_tolerance": 3,
#     # "edge_min_length": 3,
#     # "min_words_vertical": 3,
#     # "min_words_horizontal": 1,
#     # "intersection_tolerance": 3,
#     # "intersection_x_tolerance": 3,
#     # "intersection_y_tolerance": 3,
#     "text_tolerance": 1,
#     "text_x_tolerance": 1,
#     "text_y_tolerance": 1,
# }


# pdf_to_xlsx('example_pdfs/FUTURA-System-Manual.pdf', 'output/FUTURA-System-Manual.xlsx', fut_table_settings)
# os.system('start excel.exe output/FUTURA-System-Manual.xlsx')

# pdf_to_xlsx('example_pdfs/HDS10M.pdf', 'output/HDS10M.xlsx')
# os.system('start excel.exe output/HDS10M.xlsx')


# POSSIBLY ADD A METHOD TO GENERATE THE TABLE VIA EXTRACTWORD FUNCTION INSTEAD FOR TRICKY TABLES
# pdf_to_xlsx('example_pdfs/ArcModbusOPC.pdf', 'output/ArcModbusOPC.xlsx', ham_table_settings, ham_file_settings)
# os.system('start excel.exe output/ArcModbusOPC.xlsx')


# pdf_to_docx('example_pdfs/HDS10M.pdf', 'output/HDS10M.docx')
# os.system('start winword.exe output/HDS10M.docx')

# pdf_to_docx('example_pdfs/LabMuffinGuideToExfoliation.pdf', 'output/LabMuffinGuideToExfoliation.docx')
# os.system('start winword.exe output/LabMuffinGuideToExfoliation.docx')





